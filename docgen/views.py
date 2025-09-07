# docgen/views.py

from datetime import datetime
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, permission_required
from accounts.access_control import is_manager
from .models import DocumentTemplate, GeneratedDocument
from .forms import DocumentTemplateForm
from docx import Document as DocxDocument
import re
import os
from django.conf import settings
from django.urls import reverse
from django.contrib import messages
from django.utils import timezone
from docx.text.run import Run
from openpyxl import load_workbook
from pdf2docx import Converter
from PyPDF2 import PdfMerger
from .forms import PDFMergeForm, PDFToWordForm
import io


# add these imports near top of docgen/views.py
import tempfile
import shutil


from django.http import FileResponse, HttpResponse



def has_manager_access(user):
    return user.has_perm('timesheet.can_approve') or is_manager(user)

@login_required
@permission_required('timesheet.can_approve')
def upload_template(request):
    if request.method == 'POST':
        form = DocumentTemplateForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('docgen:list-templates')
    else:
        form = DocumentTemplateForm()
    return render(request, 'docgen/upload_template.html', {'form': form})


@login_required
@permission_required('timesheet.can_approve')
def list_templates(request):
    templates = DocumentTemplate.objects.all()
    return render(request, 'docgen/template_list.html', {'templates': templates})



def extract_placeholders(docx_path):
    doc = DocxDocument(docx_path)
    placeholders = set()

    # Check paragraphs
    for para in doc.paragraphs:
        placeholders.update(re.findall(r'\{\{(.*?)\}\}', para.text))

    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                placeholders.update(re.findall(r'\{\{(.*?)\}\}', cell.text))

    return sorted(placeholders)





def replace_text_in_runs(runs, replacements):
    for run in runs:
        for key, val in replacements.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, val)







from django.contrib import messages

@login_required
@permission_required('timesheet.can_approve')
def delete_template(request, template_id):
    template = get_object_or_404(DocumentTemplate, id=template_id)
    if request.method == 'POST':
        template.delete()
        messages.success(request, "Template deleted successfully.")
        return redirect('docgen:list-templates')
    return render(request, 'docgen/confirm_delete.html', {'template': template})
    
def update_excel_file(input_values, excel_path):
    wb = load_workbook(excel_path)
    ws = wb.active
    ws['D4'] = input_values.get('Old_Gross', 0)
    ws['E4'] = input_values.get('New_Gross', 0)
    ws['C1'] = input_values.get('Emp_Id', '')
    ws['C2'] = input_values.get('Emp_Name', '')
    ws['C3'] = input_values.get('Emp_Grade', '')
    wb.save(excel_path)

@login_required
@permission_required('timesheet.can_approve')
def fill_template(request, template_id):
    template = get_object_or_404(DocumentTemplate, id=template_id)
    docx_path = template.template_file.path

    # Identify if template needs Excel inputs
    needs_excel = 'revision' in template.name.lower()

    placeholders = extract_placeholders(docx_path)

    if request.method == 'POST':
        filled_data = {key: request.POST.get(key, '') for key in placeholders}

        if needs_excel:
            # Collect Excel fields only if needed
            Old_Gross = float(request.POST.get('Old_Gross', 0))
            New_Gross = float(request.POST.get('New_Gross', 0))
            Emp_Id = request.POST.get('Emp_Id', '')
            Emp_Name = request.POST.get('Emp_Name', '')
            Emp_Grade = request.POST.get('Emp_Grade', '')

            excel_template_path = os.path.join(settings.MEDIA_ROOT, 'excel_templates', 'salary_template.xlsx')
            update_excel_file({
                'Old_Gross': Old_Gross,
                'New_Gross': New_Gross,
                'Emp_Id': Emp_Id,
                'Emp_Name': Emp_Name,
                'Emp_Grade': Emp_Grade,
            }, excel_template_path)

        # Replace placeholders in Word
        doc = DocxDocument(docx_path)
        for paragraph in doc.paragraphs:
            for key, value in filled_data.items():
                paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", value)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in filled_data.items():
                        cell.text = cell.text.replace(f"{{{{{key}}}}}", value)

        # Save the document
        timestamp = timezone.now().strftime('%Y%m%d%H%M%S')
        filename = f"{template.name.replace(' ', '_')}_{timestamp}_filled.docx"
        output_path = os.path.join(settings.MEDIA_ROOT, 'generated_docs', filename)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)

        GeneratedDocument.objects.create(
            template=template,
            filled_data=filled_data,
            generated_docx=f'generated_docs/{filename}'
        )

        return HttpResponse(f"""
            <div style='padding:20px; font-family:Arial;'>
                <h3>Document generated successfully: {filename}</h3>
                <p><a href='/media/generated_docs/{filename}' target='_blank'>📄 Download Word File</a></p>
            </div>
        """)

    return render(request, 'docgen/fill_template.html', {
        'template': template,
        'placeholders': placeholders,
        'needs_excel': needs_excel
    })
   







@login_required
@permission_required('timesheet.can_approve')
def combined_pdf_tools_view(request):
    merge_form = PDFMergeForm()
    convert_form = PDFToWordForm()
    error_message = None

    if request.method == 'POST':
        # -------------------- PDF Merge --------------------
        if 'merge_submit' in request.POST:
            merge_form = PDFMergeForm(request.POST, request.FILES)
            if merge_form.is_valid():
                # Expecting fields named 'pdf1' and 'pdf2' in the form
                f1 = request.FILES.get('pdf1')
                f2 = request.FILES.get('pdf2')

                if not f1 or not f2:
                    error_message = "Please upload both PDF files to merge."
                else:
                    # Use PdfMerger directly with file-like objects when possible
                    try:
                        merger = PdfMerger()
                        # PyPDF2 PdfMerger accepts file-like objects; ensure pointer at 0
                        f1.seek(0)
                        f2.seek(0)
                        merger.append(f1)
                        merger.append(f2)
                        out = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
                        try:
                            merger.write(out)
                            merger.close()
                            out.flush()
                            out.seek(0)
                            resp = FileResponse(open(out.name, "rb"), as_attachment=True, filename="merged.pdf")
                            return resp
                        finally:
                            out.close()
                    except Exception as e:
                        error_message = f"Failed to merge PDFs: {e}"

        # -------------------- PDF -> Word --------------------
        elif 'convert_submit' in request.POST:
            convert_form = PDFToWordForm(request.POST, request.FILES)
            if convert_form.is_valid():
                uploaded_pdf = request.FILES.get('pdf_file')
                if not uploaded_pdf:
                    error_message = "Please select a PDF file to convert."
                elif not uploaded_pdf.name.lower().endswith('.pdf'):
                    error_message = "Please upload a valid PDF file."
                else:
                    # Write uploaded PDF to temp file and convert using pdf2docx
                    in_tmp = None
                    out_tmp = None
                    try:
                        in_tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
                        in_tmp.write(uploaded_pdf.read())
                        in_tmp.flush()
                        in_tmp.close()

                        out_tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                        out_tmp.close()  # pdf2docx will create/overwrite file

                        cv = Converter(in_tmp.name)
                        try:
                            cv.convert(out_tmp.name, start=0, end=None)
                        finally:
                            cv.close()

                        # return docx as FileResponse
                        resp = FileResponse(open(out_tmp.name, "rb"), as_attachment=True, filename="converted.docx")
                        return resp

                    except Exception as e:
                        error_message = f"Conversion failed: {e}"
                    finally:
                        # cleanup temp files (if they exist)
                        try:
                            if in_tmp and os.path.exists(in_tmp.name):
                                os.remove(in_tmp.name)
                        except Exception:
                            pass
                        # Note: do not remove out_tmp before returning file — handled by OS after close
                        # We'll attempt to remove it as well (safe if file still closed)
                        try:
                            if out_tmp and os.path.exists(out_tmp.name):
                                os.remove(out_tmp.name)
                        except Exception:
                            pass

    context = {
        'merge_form': merge_form,
        'convert_form': convert_form,
        'error_message': error_message,
    }
    return render(request, 'docgen/combined_pdf_tools.html', context)

